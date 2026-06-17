from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        if not item:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(str(item))


def add_kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value or "")


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def write_brd(input_data: dict, output_path: Path) -> None:
    doc = Document()
    style_document(doc)

    overview = input_data.get("overview", {})
    summary = input_data.get("summary", {})
    targeting = input_data.get("targeting", {})
    delivery = input_data.get("delivery", {})
    governance = input_data.get("governance", {})

    doc.add_paragraph("Business Requirements Document", style="Title")
    add_kv(doc, "Campaign name", str(overview.get("campaign_name", "")))
    add_kv(doc, "Internal code", str(overview.get("internal_campaign_code", "")))
    add_kv(doc, "Owner", str(governance.get("owner", "")))
    add_kv(doc, "Date", str(governance.get("document_date", "")))
    add_kv(doc, "Version", str(governance.get("current_version", "")))

    add_section_heading(doc, "1. Business Context")
    add_kv(doc, "What is happening", str(summary.get("business_context", "")))
    add_kv(doc, "Why this campaign is needed now", str(summary.get("why_now", "")))
    add_kv(doc, "Related initiative", str(summary.get("related_initiative", "")))

    add_section_heading(doc, "2. Objectives and Success Measures")
    add_kv(doc, "Primary objective", str(summary.get("primary_objective", "")))
    add_kv(doc, "Secondary objective", str(summary.get("secondary_objective", "")))
    add_kv(doc, "Hypothesis", str(summary.get("hypothesis", "")))
    if summary.get("kpis"):
        add_section_heading(doc, "KPIs", 2)
        add_bullets(doc, [str(item) for item in summary.get("kpis", [])])

    add_section_heading(doc, "3. Audience and Eligibility")
    add_kv(doc, "Primary audience", str(targeting.get("audience_description", "")))
    for title, key in [
        ("Inclusion criteria", "inclusions"),
        ("Exclusion criteria", "exclusions"),
        ("Suppression rules", "suppressions"),
    ]:
        if targeting.get(key):
            add_section_heading(doc, title, 2)
            add_bullets(doc, [str(item) for item in targeting.get(key, [])])
    add_kv(
        doc,
        "Control group approach",
        str(targeting.get("control_group", {}).get("description", "")),
    )

    add_section_heading(doc, "4. Journey and Delivery Logic")
    if overview.get("channels"):
        add_section_heading(doc, "Channels in scope", 2)
        add_bullets(doc, [str(item) for item in overview.get("channels", [])])
    add_kv(doc, "Timing rules", str(targeting.get("timing_rules", "")))
    add_kv(doc, "Trigger or schedule logic", str(targeting.get("trigger_logic", "")))
    add_kv(doc, "Operational rules", str(targeting.get("operational_rules", "")))
    if targeting.get("deliveries"):
        add_section_heading(doc, "Delivery sequence", 2)
        for item in targeting.get("deliveries", []):
            p = doc.add_paragraph(style="List Bullet")
            label = item.get("label", "")
            channel = item.get("channel", "")
            launch_date = item.get("launch_date", "")
            p.add_run(f"{label}").bold = True
            suffix = " | ".join(part for part in [channel, launch_date] if part)
            if suffix:
                p.add_run(f" - {suffix}")

    add_section_heading(doc, "5. Messaging and Offer Strategy")
    add_kv(doc, "Core proposition", str(summary.get("single_minded_message", "")))
    add_kv(doc, "Offer details", str(summary.get("offer", "")))
    add_kv(doc, "Primary CTA", str(summary.get("primary_cta", "")))
    add_kv(doc, "Tone and personalization approach", str(delivery.get("tone_and_personalization", "")))

    add_section_heading(doc, "6. Content and Creative Requirements")
    add_kv(doc, "Subject line direction", str(delivery.get("subject_line", "")))
    add_kv(doc, "Content summary", str(delivery.get("content_summary", "")))
    if delivery.get("asset_requirements"):
        add_section_heading(doc, "Required assets", 2)
        add_bullets(doc, [str(item) for item in delivery.get("asset_requirements", [])])
    if delivery.get("content_modules"):
        add_section_heading(doc, "Content modules", 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Module"
        hdr[1].text = "Audience"
        hdr[2].text = "Rule"
        hdr[3].text = "Notes"
        for item in delivery.get("content_modules", []):
            row = table.add_row().cells
            row[0].text = str(item.get("module", ""))
            row[1].text = str(item.get("audience", ""))
            row[2].text = str(item.get("rule", ""))
            row[3].text = str(item.get("notes", ""))

    add_section_heading(doc, "7. Data and Personalization Requirements")
    if delivery.get("personalization_fields"):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Type"
        hdr[1].text = "Description"
        hdr[2].text = "Field name"
        hdr[3].text = "Default / skip condition"
        for item in delivery.get("personalization_fields", []):
            row = table.add_row().cells
            row[0].text = str(item.get("type", ""))
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("field_name", ""))
            row[3].text = str(item.get("default_value_or_skip_condition", ""))

    add_section_heading(doc, "8. Testing, Risks, and Dependencies")
    if delivery.get("qa_checkpoints"):
        add_section_heading(doc, "QA checkpoints", 2)
        add_bullets(doc, [str(item) for item in delivery.get("qa_checkpoints", [])])
    if summary.get("dependencies"):
        add_section_heading(doc, "Dependencies", 2)
        add_bullets(doc, [str(item) for item in summary.get("dependencies", [])])
    if summary.get("risks"):
        add_section_heading(doc, "Risks", 2)
        add_bullets(doc, [str(item) for item in summary.get("risks", [])])
    if summary.get("compliance_considerations"):
        add_section_heading(doc, "Compliance considerations", 2)
        add_bullets(doc, [str(item) for item in summary.get("compliance_considerations", [])])

    add_section_heading(doc, "9. Approvals, Assumptions, and Open Questions")
    if governance.get("approvals"):
        add_section_heading(doc, "Required approvers", 2)
        add_bullets(doc, [str(item.get("name", "")) for item in governance.get("approvals", [])])
    if summary.get("assumptions"):
        add_section_heading(doc, "Assumptions", 2)
        add_bullets(doc, [str(item) for item in summary.get("assumptions", [])])
    if summary.get("open_questions"):
        add_section_heading(doc, "Open questions", 2)
        add_bullets(doc, [str(item) for item in summary.get("open_questions", [])])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    input_data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    write_brd(input_data, Path(args.output))


if __name__ == "__main__":
    main()
